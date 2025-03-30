import os
import docx
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt, RGBColor, Cm
from docx2pdf import convert
import io
import matplotlib.pyplot as plt
from PIL import Image
from docx import Document
from docx.shared import Inches

content_dict_s = {
    'logo_dir': '',

    'Name': '',
    'Address': '',
    'Tel': '',
    'Website': '',
    'E-mail': '',

    'Name/order': '',
    'Data/time': '',
    # gmm
    'Norm[gmm]': '',
    'Radius of balancing': '',
    # g
    'Norm class[G]': '',
    'Mass[kg]': '',
    'Nominal RPM': '',

    'result_dir': 'result2.png',

    'result': 'Rotor is unbalanced',
    'Cal RPM': '3600',
    'Test result': 'unbalanced',
    'Correction mass[g]/[°]': '',
    'Residual unbalance[gmm]': ''
}
content_dict_d = {
    'logo_dir': "big.png",

    'Name': 'LGT',
    'Address': 'Shanghai',
    'Website': 'www.baidu.com',
    'Tel': '1234567890',
    'E-mail': 'xxx@qq.com',

    'Name/order': 'Test/20241535',
    'Data/time': '2024-08-11/22:23',
    'Norm[gmm]': ('0.12', '0.12'),
    'Radius of balancing': ('10', '10'),
    'Norm class[G]': 'gmm',
    'Mass[kg]': '0.5',
    'Nominal RPM': '1500',

    'result_dir': "result1.png",

    'result': 'Rotor is unbalanced',
    'Cal RPM': '3600',
    'Test result': ('balanced', 'unbalanced'),
    'Correction mass[g]/[°]': ('0.01/271', '0.015/230'),
    'Residual unbalance[gmm]': ('0.1', '0.15')
}


def docx_s(content_dict):
    doc = docx.Document()
    # 设置纸张大小为 A4
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    # 设置页边距为A4默认页边距的一半
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.626)
    section.right_margin = Inches(0.626)
    # 标题
    heading = doc.add_heading('MEASUREMENT PROTOCOL')
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(16)
    for run in heading.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)
    heading.alignment = 1

    # 第一个表格结构
    table_company = doc.add_table(rows=5, cols=3)
    start_cell = table_company.cell(0, 0)
    end_cell = table_company.cell(4, 0)
    start_cell.merge(end_cell)
    # logo和表头
    para_logo = start_cell.paragraphs[0]
    para_logo.add_run().add_picture(content_dict['logo_dir'], width=Inches(2))
    start_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para_logo.alignment = 1
    para_logo.paragraph_format.space_before = 0
    para_logo.paragraph_format.space_after = 0
    table_company.cell(0, 1).text = 'Name'
    table_company.cell(1, 1).text = 'Address'
    table_company.cell(2, 1).text = 'Website'
    table_company.cell(3, 1).text = 'Tel'
    table_company.cell(4, 1).text = 'E-mail'
    # 第一个表格内容
    table_company.cell(0, 2).text = content_dict['Name']
    table_company.cell(1, 2).text = content_dict['Address']
    table_company.cell(2, 2).text = content_dict['Website']
    table_company.cell(3, 2).text = content_dict['Tel']
    table_company.cell(4, 2).text = content_dict['E-mail']
    # 第一个表格单元格中的字体、颜色
    for row in table_company.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = 1
                for run in paragraph.runs:
                    # run.font.color.rgb = RGBColor(91, 155, 213)
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(12)  # 设置字体大小

    # 空行
    para_space = doc.add_paragraph()
    para_space.paragraph_format.space_before = Pt(0)
    para_space.paragraph_format.space_after = Pt(0)

    # 第二个表格表头
    table_rotors = doc.add_table(rows=8, cols=2)
    table_rotors.cell(0, 0).text = 'Rotor'
    table_rotors.cell(0, 1).text = 'First correction plane'
    table_rotors.cell(1, 0).text = 'Name/order'
    table_rotors.cell(2, 0).text = 'Data/time'
    table_rotors.cell(3, 0).text = 'Norm[gmm]'
    table_rotors.cell(4, 0).text = 'Radius of balancing'
    table_rotors.cell(5, 0).text = 'Norm class[G]'
    table_rotors.cell(6, 0).text = 'Mass[kg]'
    table_rotors.cell(7, 0).text = 'Nominal RPM'
    # 第二个表格的内容
    table_rotors.cell(1, 1).text = content_dict['Name/order']
    table_rotors.cell(2, 1).text = content_dict['Data/time']
    table_rotors.cell(3, 1).text = content_dict['Norm[gmm]']
    table_rotors.cell(4, 1).text = content_dict['Radius of balancing']
    table_rotors.cell(5, 1).text = content_dict['Norm class[G]']
    table_rotors.cell(6, 1).text = content_dict['Mass[kg]']
    table_rotors.cell(7, 1).text = content_dict['Nominal RPM']
    # 第二个表格单元格中文字的字体、对齐方式
    for row in table_rotors.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = 1
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(12)

    # 空行
    para_space = doc.add_paragraph()
    para_space.paragraph_format.space_before = Pt(0)
    para_space.paragraph_format.space_after = Pt(0)

    # 结果图
    table_result_img = doc.add_table(rows=1, cols=1)
    single_cell = table_result_img.cell(0, 0)
    para_1 = single_cell.paragraphs[0]
    para_1.add_run().add_picture(content_dict['result_dir'], height=Inches(3))
    # para_1.add_run().add_picture(content_dict['result_dir'])
    para_1.alignment = 1
    # 结果
    para_result = doc.add_paragraph(content_dict['Test result'])
    para_result.paragraph_format.space_before = Pt(0)
    para_result.paragraph_format.space_after = Pt(0)
    para_result.alignment = 1
    for run in para_result.runs:
        run.font.name = '微软雅黑'
        run.font.bold = True

    # 第四个表格表头
    table_result = doc.add_table(5, 2)
    table_result.cell(0, 0).text = 'Actual result'
    table_result.cell(0, 1).text = 'First correction plane'
    table_result.cell(1, 0).text = 'Cal RPM'
    table_result.cell(2, 0).text = 'Test result'
    table_result.cell(3, 0).text = 'Correction mass[g]/[°]'
    table_result.cell(4, 0).text = 'Residual unbalance[gmm]'
    # 第四个表格内容
    table_result.cell(1, 1).text = content_dict['Cal RPM']
    table_result.cell(2, 1).text = content_dict['Test result_1']
    table_result.cell(3, 1).text = content_dict['Correction mass[g]/[°]']
    table_result.cell(4, 1).text = content_dict['Residual unbalance[gmm]']
    # 第四个表格单元格中文字的字体、对齐方式
    for row in table_result.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = 1
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(12)

    # 设置表格样式
    table_company.style = 'Table Grid'
    table_rotors.style = 'Table Grid'
    table_result_img.style = 'Table Grid'
    table_result.style = 'Table Grid'

    standard = "gmm" if content_dict['Norm class[G]'] == "" else "G"
    doc.save(f"MEASUREMENT PROTOCOL-1P-{standard}.docx")


def docx_d(content_dict):
    doc = docx.Document()
    # 设置纸张大小为 A4
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    # 设置页边距为A4默认页边距的一半
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.626)
    section.right_margin = Inches(0.626)
    # 标题
    heading = doc.add_heading('MEASUREMENT PROTOCOL')
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(16)
    for run in heading.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)
    heading.alignment = 1

    # 第一个表格结构
    table_company = doc.add_table(rows=5, cols=3)
    start_cell = table_company.cell(0, 0)
    end_cell = table_company.cell(4, 0)
    start_cell.merge(end_cell)
    # logo和表头
    para_logo = start_cell.paragraphs[1]
    para_logo.add_run().add_picture(content_dict['logo_dir'], width=Inches(2))
    start_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para_logo.alignment = 1
    para_logo.paragraph_format.space_before = 0
    para_logo.paragraph_format.space_after = 0
    table_company.cell(0, 1).text = 'Name'
    table_company.cell(1, 1).text = 'Address'
    table_company.cell(2, 1).text = 'Website'
    table_company.cell(3, 1).text = 'Tel'
    table_company.cell(4, 1).text = 'E-mail'
    # 第一个表格内容
    table_company.cell(0, 2).text = content_dict['Name']
    table_company.cell(1, 2).text = content_dict['Address']
    table_company.cell(2, 2).text = content_dict['Website']
    table_company.cell(3, 2).text = content_dict['Tel']
    table_company.cell(4, 2).text = content_dict['E-mail']
    # 第一个表格单元格中文字的对齐方式、字体、颜色
    for row in table_company.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = 1
                for run in paragraph.runs:
                    # run.font.color.rgb = RGBColor(91, 155, 213)
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(12)

    # 空行
    para_space = doc.add_paragraph()
    para_space.paragraph_format.space_before = Pt(0)
    para_space.paragraph_format.space_after = Pt(0)

    # 第二个表格表结构
    table_rotors = doc.add_table(rows=8, cols=3)
    start_cell = table_rotors.cell(1, 1)
    end_cell = table_rotors.cell(1, 2)
    start_cell.merge(end_cell)
    start_cell = table_rotors.cell(2, 1)
    end_cell = table_rotors.cell(2, 2)
    start_cell.merge(end_cell)
    start_cell = table_rotors.cell(5, 1)
    end_cell = table_rotors.cell(5, 2)
    start_cell.merge(end_cell)
    start_cell = table_rotors.cell(6, 1)
    end_cell = table_rotors.cell(6, 2)
    start_cell.merge(end_cell)
    start_cell = table_rotors.cell(7, 1)
    end_cell = table_rotors.cell(7, 2)
    start_cell.merge(end_cell)
    # 第二个表格表头
    table_rotors.cell(0, 0).text = 'Rotor'
    table_rotors.cell(0, 1).text = 'First correction plane'
    table_rotors.cell(0, 2).text = 'Second correction plane'
    table_rotors.cell(1, 0).text = 'Name/order'
    table_rotors.cell(2, 0).text = 'Data/time'
    table_rotors.cell(3, 0).text = 'Norm[gmm]'
    table_rotors.cell(4, 0).text = 'Radius of balancing'
    table_rotors.cell(5, 0).text = 'Norm class[G]'
    table_rotors.cell(6, 0).text = 'Mass[kg]'
    table_rotors.cell(7, 0).text = 'Nominal RPM'

    # 第二个表格内容
    table_rotors.cell(1, 1).text = content_dict['Name/order']
    table_rotors.cell(2, 1).text = content_dict['Data/time']
    table_rotors.cell(3, 1).text = content_dict['Norm[gmm]_1']
    table_rotors.cell(3, 2).text = content_dict['Norm[gmm]_2']
    table_rotors.cell(4, 1).text = content_dict['Radius of balancing_1']
    table_rotors.cell(4, 2).text = content_dict['Radius of balancing_2']
    table_rotors.cell(5, 1).text = content_dict['Norm class[G]']
    table_rotors.cell(6, 1).text = content_dict['Mass[kg]']
    table_rotors.cell(7, 1).text = content_dict['Nominal RPM']
    # 第二个表格单元格中文字的对齐方式、字体、颜色
    for row in table_rotors.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = 1
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(12)

    # 空行
    para_space = doc.add_paragraph()
    para_space.paragraph_format.space_before = Pt(0)
    para_space.paragraph_format.space_after = Pt(0)

    # 结果图
    table_result_img = doc.add_table(rows=1, cols=1)
    single_cell = table_result_img.cell(0, 0)
    para_1 = single_cell.paragraphs[0]
    para_1.add_run().add_picture(content_dict['result_dir'], height=Inches(3))
    para_1.alignment = 1
    # 结果
    para_result = doc.add_paragraph(content_dict['Test result'])
    para_result.paragraph_format.space_before = Pt(0)
    para_result.paragraph_format.space_after = Pt(0)
    para_result.alignment = 1
    for run in para_result.runs:
        run.font.name = '微软雅黑'
        run.font.bold = True

    # 第四个表格表结构
    table_result = doc.add_table(5, 3)
    start_cell = table_result.cell(1, 1)
    end_cell = table_result.cell(1, 2)
    start_cell.merge(end_cell)
    # 第四个表格表头
    table_result.cell(0, 0).text = 'Actual result'
    table_result.cell(0, 1).text = 'First correction plane'
    table_result.cell(0, 2).text = 'Second correction plane'
    table_result.cell(1, 0).text = 'Cal RPM'
    table_result.cell(2, 0).text = 'Test result'
    table_result.cell(3, 0).text = 'Correction mass[g]/[°]'
    table_result.cell(4, 0).text = 'Residual unbalance[gmm]'

    # 第四个表格内容
    table_result.cell(1, 1).text = content_dict['Cal RPM']
    table_result.cell(2, 1).text = content_dict['Test result_1']
    table_result.cell(2, 2).text = content_dict['Test result_2']
    table_result.cell(3, 1).text = content_dict['Correction mass[g]/[°]_1']
    table_result.cell(3, 2).text = content_dict['Correction mass[g]/[°]_2']
    table_result.cell(4, 1).text = content_dict['Residual unbalance[gmm]_1']
    table_result.cell(4, 2).text = content_dict['Residual unbalance[gmm]_2']
    # 第四个表格单元格中文字的字体、对齐方式
    for row in table_result.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = 1
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(12)

    # 设置表格样式
    table_company.style = 'Table Grid'
    table_rotors.style = 'Table Grid'
    table_result_img.style = 'Table Grid'
    table_result.style = 'Table Grid'

    standard = "gmm" if content_dict['Norm class[G]'] == "" else "G"
    doc.save(f"MEASUREMENT PROTOCOL-2P-{standard}.docx")


def docx_ds(content_dict):
    doc = docx.Document()
    # 设置纸张大小为 A4
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    # 设置页边距为A4默认页边距的一半
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.626)
    section.right_margin = Inches(0.626)
    # 表头
    heading = doc.add_heading('MEASUREMENT PROTOCOL')
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(16)
    for run in heading.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)
    heading.alignment = 1

    # 第一个表格结构
    table_company = doc.add_table(rows=5, cols=3)
    start_cell = table_company.cell(0, 0)
    end_cell = table_company.cell(4, 0)
    start_cell.merge(end_cell)
    # logo和表头
    para_logo = start_cell.paragraphs[0]
    para_logo.add_run().add_picture(content_dict['logo_dir'], width=Inches(2))
    para_logo.alignment = 1
    start_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para_logo.paragraph_format.space_before = 0
    para_logo.paragraph_format.space_after = 0
    table_company.cell(0, 1).text = 'Name'
    table_company.cell(1, 1).text = 'Address'
    table_company.cell(2, 1).text = 'Website'
    table_company.cell(3, 1).text = 'Tel'
    table_company.cell(4, 1).text = 'E-mail'
    # 第一个表格内容
    table_company.cell(0, 2).text = content_dict['Name']
    table_company.cell(1, 2).text = content_dict['Address']
    table_company.cell(2, 2).text = content_dict['Website']
    table_company.cell(3, 2).text = content_dict['Tel']
    table_company.cell(4, 2).text = content_dict['E-mail']
    # 第一个表格单元格中文字的对齐方式、字体、颜色
    for row in table_company.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = 1
                for run in paragraph.runs:
                    # run.font.color.rgb = RGBColor(91, 155, 213)
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(12)

    # 空行
    para_space = doc.add_paragraph()
    para_space.paragraph_format.space_before = Pt(0)
    para_space.paragraph_format.space_after = Pt(0)

    # 第二个表格表头
    table_rotors = doc.add_table(rows=8, cols=2)
    table_rotors.cell(0, 0).text = 'Rotor'
    table_rotors.cell(0, 1).text = 'Couple plane'
    table_rotors.cell(1, 0).text = 'Name/order'
    table_rotors.cell(2, 0).text = 'Data/time'
    table_rotors.cell(3, 0).text = 'Norm[gmm]'
    table_rotors.cell(4, 0).text = 'Radius of balancing'
    table_rotors.cell(5, 0).text = 'Norm class[G]'
    table_rotors.cell(6, 0).text = 'Mass[kg]'
    table_rotors.cell(7, 0).text = 'Nominal RPM'
    # 第二个表格的内容
    table_rotors.cell(1, 1).text = content_dict['Name/order']
    table_rotors.cell(2, 1).text = content_dict['Data/time']
    table_rotors.cell(3, 1).text = content_dict['Norm[gmm]']
    table_rotors.cell(4, 1).text = content_dict['Radius of balancing']
    table_rotors.cell(5, 1).text = content_dict['Norm class[G]']
    table_rotors.cell(6, 1).text = content_dict['Mass[kg]']
    table_rotors.cell(7, 1).text = content_dict['Nominal RPM']
    # 第二个表格单元格中文字的字体、对齐方式
    for row in table_rotors.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = 1
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(12)

    # 空行
    para_space = doc.add_paragraph()
    para_space.paragraph_format.space_before = Pt(0)
    para_space.paragraph_format.space_after = Pt(0)

    # 结果图
    table_result_img = doc.add_table(rows=1, cols=1)
    single_cell = table_result_img.cell(0, 0)
    para_space = single_cell.add_paragraph()
    para_space = single_cell.add_paragraph()
    para_space = single_cell.add_paragraph()
    para_1 = single_cell.paragraphs[3]
    para_space = single_cell.add_paragraph()
    para_space = single_cell.add_paragraph()
    para_1.add_run().add_picture(content_dict['result_dir'], width=Inches(6.88))
    # para_1.alignment = 1
    # 结果
    para_result = doc.add_paragraph(content_dict['Test result'])
    para_result.paragraph_format.space_before = Pt(0)
    para_result.paragraph_format.space_after = Pt(0)
    para_result.alignment = 1
    for run in para_result.runs:
        run.font.name = '微软雅黑'
        run.font.bold = True

    # 第四个表格表头
    table_result = doc.add_table(5, 2)
    table_result.cell(0, 0).text = 'Actual result'
    table_result.cell(0, 1).text = 'Couple correction plane'
    table_result.cell(1, 0).text = 'Cal RPM'
    table_result.cell(2, 0).text = 'Test result'
    table_result.cell(3, 0).text = 'Correction mass[g]/[°]'
    table_result.cell(4, 0).text = 'Residual unbalance[gmm]'
    # 第四个表格内容
    table_result.cell(1, 1).text = content_dict['Cal RPM']
    table_result.cell(2, 1).text = content_dict['Test result_3']
    table_result.cell(3, 1).text = content_dict['Correction mass[g]/[°]']
    table_result.cell(4, 1).text = content_dict['Residual unbalance[gmm]']
    # 第四个表格单元格中文字的字体、对齐方式
    for row in table_result.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = 1
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(12)

    # 设置表格样式
    table_company.style = 'Table Grid'
    table_rotors.style = 'Table Grid'
    table_result_img.style = 'Table Grid'
    table_result.style = 'Table Grid'
    standard = "gmm" if content_dict['Norm class[G]'] == "" else "G"
    doc.save(f"MEASUREMENT PROTOCOL-couple.docx")


def docx_to_pdf_and_delete_docx():
    convert("D:/项目文件/Python/docx-pdf-test/Example_onepale.docx",
            "D:/项目文件/Python/docx-pdf-test/Example_onepale.pdf")




