import os
from pathlib import Path

file_run_root = Path(__file__).parent

os.chdir(file_run_root)

import docx
from docx2pdf import convert
import os

def create_docx_2(content_dict):
    # double
    doc = docx.Document()
    heading = doc.add_heading('MEASUREMENT PROTOCOL')
    heading.alignment = 1 # center
    
    pic_logo = doc.add_picture(content_dict['logo_dir'])

    table_company = doc.add_table(rows=5, cols=2)
    table_company.cell(0, 0).text = 'Name'
    table_company.cell(1, 0).text = 'Address'
    table_company.cell(2, 0).text = 'Website'
    table_company.cell(3, 0).text = 'Tel'
    table_company.cell(4, 0).text = 'E-mail'
    
    table_company.cell(0, 1).text = content_dict['Name']
    table_company.cell(1, 1).text = content_dict['Address']
    table_company.cell(2, 1).text = content_dict['Website']
    table_company.cell(3, 1).text = content_dict['Tel']
    table_company.cell(4, 1).text = content_dict['E-mail']
    
    para_space = doc.add_paragraph('')
    
    table_rotors = doc.add_table(rows=8, cols = 3)
    table_rotors.cell(0, 0).text = 'Rotor'
    table_rotors.cell(0, 1).text = 'First correction plane'
    table_rotors.cell(0, 2).text = 'Second correction plane'
    
    table_rotors.cell(1, 0).text = 'Name/order'
    table_rotors.cell(2, 0).text = 'Data/time'
    table_rotors.cell(3, 0).text = 'Norm[gmm]'
    table_rotors.cell(4, 0).text = 'Radius of balancing'
    table_rotors.cell(5, 0).text = 'Norm class[G]'
    table_rotors.cell(6, 0).text = 'Mass[kg]'
    table_rotors.cell(7, 0).text = 'Nominal RPM'
    
    table_rotors.cell(1, 1).merge(table_rotors.cell(1, 2)).text = content_dict['Name/order']
    table_rotors.cell(2, 1).merge(table_rotors.cell(2, 2)).text = content_dict['Data/time']
    table_rotors.cell(3, 1).text = content_dict['Norm[gmm]'][0]
    table_rotors.cell(3, 2).text = content_dict['Norm[gmm]'][1]

    table_rotors.cell(4, 1).text = content_dict['Radius of balancing'][0]
    table_rotors.cell(4, 2).text = content_dict['Radius of balancing'][1]
    
    table_rotors.cell(5, 1).merge(table_rotors.cell(5, 2)).text = content_dict['Norm class[G]']
    table_rotors.cell(6, 1).merge(table_rotors.cell(6, 2)).text = content_dict['Mass[kg]']
    table_rotors.cell(7, 1).merge(table_rotors.cell(7, 2)).text = content_dict['Nominal RPM']
    
    
    pic_result = doc.add_picture(content_dict['result_dir'])
    
    para_result = doc.add_paragraph(content_dict['result'])
    para_result.alignment = 1 # center
    
    table_result = doc.add_table(5, 3)
    table_result.cell(0, 0).text = 'Actual result'
    table_result.cell(0, 1).text = 'First correction plane'
    table_result.cell(0, 2).text = 'Second correction plane'
    
    table_result.cell(1, 0).text = 'Cal RPM'
    table_result.cell(2, 0).text = 'Test result'
    table_result.cell(3, 0).text = 'Correction mass[g]/[°]'
    table_result.cell(4, 0).text = 'Residual unbalance[gmm]'
    table_result.cell(1, 1).merge(table_result.cell(1, 2)).text = content_dict['Cal RPM']
    table_result.cell(2, 1).text = content_dict['Test result'][0]
    table_result.cell(2, 2).text = content_dict['Test result'][1]
    table_result.cell(3, 1).text = content_dict['Correction mass[g]/[°]'][0]
    table_result.cell(3, 2).text = content_dict['Correction mass[g]/[°]'][1]
    table_result.cell(4, 1).text = content_dict['Residual unbalance[gmm]'][0]
    table_result.cell(4, 2).text = content_dict['Residual unbalance[gmm]'][1]
    
    
    table_company.style = 'Table Grid'
    table_rotors.style = 'Table Grid'
    table_result.style = 'Table Grid'

    # 保存文档
    doc.save("MEASUREMENT_PROTOCOL2.docx")
    
def create_docx_1(content_dict):
    # single
    doc = docx.Document()
    heading = doc.add_heading('MEASUREMENT PROTOCOL')
    heading.alignment = 1 # center
    
    pic_logo = doc.add_picture(content_dict['logo_dir'])

    table_company = doc.add_table(rows=5, cols=2)
    table_company.cell(0, 0).text = 'Name'
    table_company.cell(1, 0).text = 'Address'
    table_company.cell(2, 0).text = 'Website'
    table_company.cell(3, 0).text = 'Tel'
    table_company.cell(4, 0).text = 'E-mail'
    
    table_company.cell(0, 1).text = content_dict['Name']
    table_company.cell(1, 1).text = content_dict['Address']
    table_company.cell(2, 1).text = content_dict['Website']
    table_company.cell(3, 1).text = content_dict['Tel']
    table_company.cell(4, 1).text = content_dict['E-mail']
    
    para_space = doc.add_paragraph('')
    
    table_rotors = doc.add_table(rows=8, cols = 2)
    table_rotors.cell(0, 0).text = 'Rotor'
    table_rotors.cell(0, 1).text = 'First correction plane'
    table_rotors.cell(0, 2).text = 'Second correction plane'
    
    table_rotors.cell(1, 0).text = 'Name/order'
    table_rotors.cell(2, 0).text = 'Data/time'
    table_rotors.cell(3, 0).text = 'Norm[gmm]'
    table_rotors.cell(4, 0).text = 'Radius of balancing'
    table_rotors.cell(5, 0).text = 'Norm class[G]'
    table_rotors.cell(6, 0).text = 'Mass[kg]'
    table_rotors.cell(7, 0).text = 'Nominal RPM'
    
    table_rotors.cell(1, 1).text = content_dict['Name/order']
    table_rotors.cell(2, 1).text = content_dict['Data/time']
    table_rotors.cell(3, 1).text = content_dict['Norm[gmm]']
    table_rotors.cell(4, 1).text = content_dict['Radius of balancing']
    table_rotors.cell(5, 1).text = content_dict['Norm class[G]']
    table_rotors.cell(6, 1).text = content_dict['Mass[kg]']
    table_rotors.cell(7, 1).text = content_dict['Nominal RPM']
    
    
    pic_result = doc.add_picture(content_dict['result_dir'])
    
    para_result = doc.add_paragraph(content_dict['result'])
    para_result.alignment = 1 # center
    
    table_result = doc.add_table(5, 2)
    table_result.cell(0, 0).text = 'Actual result'
    table_result.cell(0, 1).text = 'First correction plane'
    
    table_result.cell(1, 0).text = 'Cal RPM'
    table_result.cell(2, 0).text = 'Test result'
    table_result.cell(3, 0).text = 'Correction mass[g]/[°]'
    table_result.cell(4, 0).text = 'Residual unbalance[gmm]'
    
    table_result.cell(1, 1).text = content_dict['Cal RPM']
    table_result.cell(2, 1).text = content_dict['Test result']
    table_result.cell(3, 1).text = content_dict['Correction mass[g]/[°]']
    table_result.cell(4, 1).text = content_dict['Residual unbalance[gmm]']
    
    
    table_company.style = 'Table Grid'
    table_rotors.style = 'Table Grid'
    table_result.style = 'Table Grid'

    # 保存文档
    doc.save("MEASUREMENT_PROTOCOL1.docx")

def convert_to_pdf():
    # 将.docx 文件转换为.pdf 文件
    convert("MEASUREMENT_PROTOCOL1.docx", file_run_root / "MEASUREMENT_PROTOCOL1.pdf")
    convert("MEASUREMENT_PROTOCOL2.docx", file_run_root / "MEASUREMENT_PROTOCOL2.pdf")

def print_pdf():
    # 自动打印.pdf 文件
    os.startfile("MEASUREMENT_PROTOCOL.pdf", "print")

if __name__ == "__main__":
    content_dict_1 = {
        'logo_dir':str(file_run_root/'img-test1.png'),
        
        'Name':'LGT',
        'Address':'Shanghai',
        'Website':'www.baidu.com',
        'Tel':'1234567890',
        'E-mail':'xxx@qq.com',
        
        'Name/order': 'Test/20241535',
        'Data/time': '2024-08-11/22:23',
        'Norm[gmm]':'0.12',
        'Radius of balancing':'10',
        'Norm class[G]': 'gmm',
        'Mass[kg]': '0.5',
        'Nominal RPM': '1500',
        
        'result_dir':str(file_run_root/'img-test2.png'),
        
        'result': 'Rotor is unbalanced',
        'Cal RPM': '3600',
        'Test result': 'unbalanced',
        'Correction mass[g]/[°]': '0.01/271',
        'Residual unbalance[gmm]': '0.1'
        
    }
    content_dict_2 = {
        'logo_dir':str(file_run_root/'img-test1.png'),
        
        'Name':'LGT',
        'Address':'Shanghai',
        'Website':'www.baidu.com',
        'Tel':'1234567890',
        'E-mail':'xxx@qq.com',
        
        'Name/order': 'Test/20241535',
        'Data/time': '2024-08-11/22:23',
        'Norm[gmm]':('0.12', '0.12'),
        'Radius of balancing':('10', '10'),
        'Norm class[G]': 'gmm',
        'Mass[kg]': '0.5',
        'Nominal RPM': '1500',
        
        'result_dir':str(file_run_root/'img-test.png'),
        
        'result': 'Rotor is unbalanced',
        'Cal RPM': '3600',
        'Test result': ('balanced', 'unbalanced'),
        'Correction mass[g]/[°]': ('0.01/271', '0.015/230'),
        'Residual unbalance[gmm]': ('0.1', '0.15')
        
    }
    create_docx_1(content_dict_1)
    create_docx_2(content_dict_2)
    convert_to_pdf()
    print_pdf()